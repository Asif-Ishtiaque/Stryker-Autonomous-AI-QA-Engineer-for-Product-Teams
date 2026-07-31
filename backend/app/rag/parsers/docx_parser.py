from __future__ import annotations

import io

from docx import Document


def parse_docx(raw: bytes, filename: str) -> list[str]:
    document = Document(io.BytesIO(raw))
    blocks: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return blocks
